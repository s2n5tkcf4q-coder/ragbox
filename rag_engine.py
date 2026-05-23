"""
rag_engine.py - RAG 引擎
负责文档处理、分块、向量化存储，以及基于知识库的问答检索生成。
"""
import os
import re
import logging
from typing import List, Dict, Tuple, Optional, Generator

import chromadb
from chromadb.config import Settings as ChromaSettings
from chromadb.utils import embedding_functions

from database import db_session
from models import SystemConfig, Document, Chunk, AuditLog
from llm_client import get_llm_client

logger = logging.getLogger(__name__)

# ChromaDB 全局客户端（懒加载）
_chroma_client = None
_collection = None


from config import Config  # 在文件顶部导入 Config 类

def _get_chroma_client() -> chromadb.Client:
    global _chroma_client
    if _chroma_client is None:
        chroma_dir = Config.CHROMA_DB_DIR
        os.makedirs(chroma_dir, exist_ok=True)
        _chroma_client = chromadb.PersistentClient(
            path=chroma_dir,
            settings=ChromaSettings(anonymized_telemetry=False)
        )
        logger.info(f"ChromaDB 客户端已初始化，路径: {chroma_dir}")
    return _chroma_client


def _get_collection() -> chromadb.Collection:
    """获取或创建向量集合（collection）"""
    global _collection
    if _collection is None:
        client = _get_chroma_client()
        # 集合名称固定，或可从配置读取
        collection_name = "knowledge_base"
        try:
            _collection = client.get_collection(collection_name)
            logger.info(f"获取已有 ChromaDB 集合: {collection_name}")
        except Exception:
            _collection = client.create_collection(
                collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            logger.info(f"创建新 ChromaDB 集合: {collection_name}")
    return _collection


def _get_embedding_function():
    """
    根据当前系统配置获取嵌入函数。
    注意：我们使用自定义的 llm_client 完成嵌入，而不是 chromadb 自带的 embedding function。
    """
    llm = get_llm_client()
    # 返回一个可调用对象，chromadb 需要 embedding_function 接受 list[str] 返回 list[list[float]]
    def embed_func(texts: List[str]) -> List[List[float]]:
        return llm.embed(texts)
    return embed_func


# -------------------- 文档分块工具 --------------------

def _split_text(text: str, chunk_size: int = Config.DEFAULT_CHUNK_SIZE,
                overlap: int = Config.DEFAULT_CHUNK_OVERLAP) -> List[str]:
    """
    递归字符分割器：按段落、句子、字符逐步切分，保证块大小大致为 chunk_size。
    """
    chunks = []
    # 按段落分割
    paragraphs = re.split(r'\n\s*\n', text)
    current_chunk = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        # 如果添加后不超过 chunk_size，直接合并
        if len(current_chunk) + len(para) <= chunk_size:
            if current_chunk:
                current_chunk += "\n\n" + para
            else:
                current_chunk = para
        else:
            # 保存当前块
            if current_chunk:
                chunks.append(current_chunk)
            # 新的段落作为新块的开始
            if len(para) > chunk_size:
                # 段落本身太长，按句子进一步拆分
                sentences = re.split(r'(?<=[。！？\.\?!])', para)
                temp_chunk = ""
                for sent in sentences:
                    if len(temp_chunk) + len(sent) <= chunk_size:
                        temp_chunk += sent
                    else:
                        if temp_chunk:
                            chunks.append(temp_chunk)
                        # 如果单句太长，粗暴截断
                        while len(sent) > chunk_size:
                            chunks.append(sent[:chunk_size])
                            sent = sent[chunk_size - overlap:]
                        temp_chunk = sent
                if temp_chunk:
                    current_chunk = temp_chunk
                else:
                    current_chunk = ""
            else:
                current_chunk = para

    if current_chunk:
        chunks.append(current_chunk)

    # 对过长的 chunk 进行二次切分（安全）
    final_chunks = []
    for ch in chunks:
        while len(ch) > chunk_size:
            final_chunks.append(ch[:chunk_size])
            ch = ch[chunk_size - overlap:]
        if ch:
            final_chunks.append(ch)
    return final_chunks


def _extract_docx_text(file_path: str) -> str:
    """从 .docx 文件提取纯文本，保留基本结构"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return '\n\n'.join(paragraphs)
    except Exception as e:
        logger.error(f"提取 docx 文本失败: {file_path}, 错误: {e}")
        raise


# -------------------- 核心处理 --------------------

def process_document(file_path: str, original_filename: str, user_id: int, file_size: int = 0) -> Document:
    """
    处理上传的文档：解析、分块、向量化并存储。
    返回 Document 模型对象。
    """
    # 检查 embedding 模型是否可用
    llm = get_llm_client()
    try:
        # 快速测试嵌入功能
        test_vec = llm.embed(["测试"])
        if not test_vec or len(test_vec[0]) == 0:
            raise RuntimeError("嵌入模型返回空向量")
    except Exception as e:
        raise RuntimeError(f"嵌入模型未加载或不可用: {e}")

    # 创建文档记录
    doc_record = Document(
        filename=os.path.basename(file_path),
        original_filename=original_filename,
        file_size=file_size,
        uploaded_by=user_id,
        status='processing'
    )
    db_session.add(doc_record)
    db_session.commit()

    try:
        # 1. 提取文本
        logger.info(f"开始处理文档: {original_filename}")
        text = _extract_docx_text(file_path)

        # 2. 分块
        chunks_text = _split_text(text)
        if not chunks_text:
            raise ValueError("文档中没有提取到有效文本")

        # 3. 向量化与存储
        collection = _get_collection()
        # 为每个块创建 ID，格式：doc_id_chunk_index
        ids = [f"doc{doc_record.id}_chunk{i}" for i in range(len(chunks_text))]
        metadatas = [{"doc_id": str(doc_record.id), "chunk_index": i} for i in range(len(chunks_text))]

        # 逐个嵌入（批量可能更快，但需确保顺序）
        embeddings = llm.embed(chunks_text)

        if not embeddings or len(embeddings) == 0:
            raise ValueError("嵌入返回空列表")
        logger.info(f"成功嵌入 {len(embeddings)} 个块，第一个向量维度: {len(embeddings[0])}")

        # 添加到 ChromaDB
        collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=chunks_text,
            metadatas=metadatas
        )

        # 4. 保存块记录到数据库
        for i, chunk_text in enumerate(chunks_text):
            chunk_record = Chunk(
                document_id=doc_record.id,
                chunk_index=i,
                content=chunk_text,
                char_count=len(chunk_text),
                embedding_id=ids[i]
            )
            db_session.add(chunk_record)

        # 更新文档状态
        doc_record.status = 'completed'
        doc_record.chunk_count = len(chunks_text)
        db_session.commit()

        # 记录审计日志
        log_entry = AuditLog(
            user_id=user_id,
            action='document_processed',
            details=f'文档 {original_filename} 处理完成，共 {len(chunks_text)} 个块'
        )
        db_session.add(log_entry)
        db_session.commit()

        logger.info(f"文档处理成功: {original_filename}, 块数: {len(chunks_text)}")
        return doc_record

    except Exception as e:
        logger.error(f"文档处理失败: {original_filename}, 错误: {e}")
        doc_record.status = 'failed'
        doc_record.error_message = str(e)
        db_session.commit()
        raise


# -------------------- 检索与问答 --------------------

def retrieve(query: str, top_n: int = None, threshold: float = None) -> List[Dict]:
    """
    检索与查询最相关的文本块
    :return: 列表，每个元素包含 'content', 'score', 'metadata'
    """
    collection = _get_collection()
    collection_meta = collection.metadata
    use_cosine = collection_meta and collection_meta.get("hnsw:space") == "cosine"

    if top_n is None:
        top_n = int(SystemConfig.get_value(db_session, 'top_n', 5))
    if threshold is None:
        threshold = float(SystemConfig.get_value(db_session, 'similarity_threshold', 0.7))

    llm = get_llm_client()
    query_embedding = llm.embed([query])[0]

    results = collection.query(
        query_embeddings=[query_embedding],
        n_results=top_n,
        include=["documents", "metadatas", "distances"]
    )

    # 整理结果
    retrieved = []
    if results['ids'] and results['ids'][0]:
        for doc, meta, dist in zip(results['documents'][0], results['metadatas'][0], results['distances'][0]):
            # ChromaDB 返回的距离是欧几里得距离或余弦距离？我们使用余弦相似度转换
            # 默认 ChromaDB 使用 L2 距离，转为相似度：相似度 = 1/(1+distance)
            if use_cosine:
                similarity = 1.0 - dist  # 余弦空间：distance 直接就是相似度
            else:
                similarity = 1.0 / (1.0 + dist)
            if similarity >= threshold:
                retrieved.append({
                    'content': doc,
                    'score': similarity,
                    'metadata': meta
                })
    # 按相似度降序
    retrieved.sort(key=lambda x: x['score'], reverse=True)
    logger.info(f"检索到 {len(retrieved)} 个相关块（阈值 {threshold}）")

    return retrieved


def generate_answer(
    query: str,
    conversation_history: List[Dict[str, str]] = None,
    stream: bool = False
) -> str | Generator[str, None, None]:
    """
    生成回答，包含思维链。
    返回格式：先返回一段完整的思考过程（【思考】...），再返回答案（【回答】...）。
    如果 stream=True，返回生成器逐步产出文本；否则返回完整字符串。
    """
    # 检索上下文
    retrieved = retrieve(query)
    context_parts = []
    for i, item in enumerate(retrieved):
        context_parts.append(f"【参考{i+1}】{item['content']}")
    context = '\n'.join(context_parts)

    # 获取系统提示词和生成参数
    system_prompt = SystemConfig.get_value(db_session, 'system_prompt',
        '你是一个专业的知识库问答助手。请严格依据提供的上下文信息回答问题，不要添加额外猜测。')
    temperature = float(SystemConfig.get_value(db_session, 'temperature', 0.7))
    top_p = float(SystemConfig.get_value(db_session, 'top_p', 0.9))
    max_tokens = int(SystemConfig.get_value(db_session, 'max_tokens', 1024))
    presence_penalty = float(SystemConfig.get_value(db_session, 'presence_penalty', 0.0))
    frequency_penalty = float(SystemConfig.get_value(db_session, 'frequency_penalty', 0.0))

    # 构建消息列表
    messages = [{"role": "system", "content": system_prompt}]
    if conversation_history:
        # 过滤掉历史中的 system 消息，只保留 user 和 assistant
        filtered_history = [
            msg for msg in conversation_history
            if msg.get('role') in ('user', 'assistant')
        ]
        # 只保留最近 10 轮对话，避免上下文过长
        messages.extend(filtered_history[-10:])

    # 用户消息：包含上下文和问题
    user_content = f"请根据以下知识库内容回答问题。\n知识库内容：\n{context}\n\n问题：{query}\n\n请先逐步思考，再给出最终答案。使用【思考】和【回答】标记分开。"
    messages.append({"role": "user", "content": user_content})

    llm = get_llm_client()
    return llm.chat(
        messages=messages,
        stream=stream,
        temperature=temperature,
        top_p=top_p,
        max_tokens=max_tokens,
        presence_penalty=presence_penalty,
        frequency_penalty=frequency_penalty,
        stop=["<|endoftext|>", "<|im_start|>"]
    )


def split_thinking_answer(text: str) -> Tuple[str, str]:
    """
    从 LLM 返回的文本中分离思维链和答案。
    预期格式：包含 '【思考】' 和 '【回答】' 标记。
    若分离失败，则将全部文本作为答案，思维链置空。
    """
    thinking = ""
    answer = text

    try:
        # 优先匹配标准格式
        think_match = re.search(r'【思考】(.*?)【回答】', text, re.DOTALL)
        if think_match:
            thinking = think_match.group(1).strip()
            answer_start = text.find('【回答】') + len('【回答】')
            answer = text[answer_start:].strip()
        elif '【思考】' in text and '【回答】' in text:
            # 顺序可能颠倒或存在额外内容，尝试分割
            parts = text.split('【回答】', 1)
            thinking = parts[0].replace('【思考】', '').strip()
            answer = parts[1].strip() if len(parts) > 1 else ''
        elif '【回答】' in text:
            # 只有回答标记，无思考
            answer = text.split('【回答】', 1)[1].strip()
        elif '【思考】' in text:
            # 只有思考标记，无回答
            thinking = text.replace('【思考】', '').strip()
            answer = ''  # 无答案
        else:
            # 完全没有标记，全部作为答案
            answer = text.strip()
            thinking = ''
    except Exception as e:
        logger.error(f"分离思考与答案时出错: {e}, 返回原始文本作为答案")
        answer = text
        thinking = ''

    # 最终保底：如果答案为空，但原文非空，则用原文作为答案
    if not answer and text.strip():
        answer = text.strip()
        thinking = thinking if thinking else '未能解析思维链'

    return thinking, answer

def is_knowledge_base_empty() -> bool:
    """检查知识库中是否有已成功处理的文档"""
    try:
        collection = _get_collection()
        return collection.count() == 0
    except Exception:
        # 如果集合不存在或出错，视为空
        return True
